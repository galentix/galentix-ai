"""
Galentix AI - Documents Router
Handles document upload and management for RAG.
"""
import os
import re
import uuid
import aiofiles
from typing import List
from datetime import datetime
from pathlib import Path
from fastapi import APIRouter, Depends, HTTPException, UploadFile, File, BackgroundTasks, Query, status
from fastapi.responses import JSONResponse
from sqlalchemy.ext.asyncio import AsyncSession
from sqlalchemy import select, func

from ..database import get_db
from ..models.document import Document
from ..models.schemas import (
    DocumentResponse,
    DocumentListResponse,
    DocumentUploadResponse,
    PaginatedResponse,
    ErrorResponse
)
from ..services.rag import get_rag_pipeline
from ..config import settings

router = APIRouter(
    prefix="/api/documents",
    tags=["documents"],
    responses={
        400: {"model": ErrorResponse, "description": "Invalid request"},
        422: {"model": ErrorResponse, "description": "Validation error"},
        429: {"model": ErrorResponse, "description": "Rate limit exceeded"},
        500: {"model": ErrorResponse, "description": "Internal server error"}
    }
)

SUPPORTED_EXTENSIONS = {".pdf", ".txt", ".md", ".docx", ".doc", ".csv", ".json"}
MAX_FILE_SIZE = 50 * 1024 * 1024
MAX_PAGE_SIZE = 100
DEFAULT_PAGE_SIZE = 20


def is_garbled_text(text: str) -> bool:
    """Detect if extracted text is garbled (spaced-out characters, etc.)."""
    if not text or len(text.strip()) < 20:
        return True
    # Check ratio of single-char words — garbled PDFs have mostly single chars
    words = text.split()
    if len(words) < 5:
        return False
    single_char_ratio = sum(1 for w in words if len(w) == 1) / len(words)
    return single_char_ratio > 0.5


def extract_pdf_text(filepath: Path) -> str:
    """Extract text from PDF using best available method.
    Fallback chain: PyMuPDF → OCR (tesseract) → pypdf."""

    text = ""

    # 1. Try PyMuPDF (best general-purpose extraction)
    try:
        import fitz  # PyMuPDF
        doc = fitz.open(str(filepath))
        text_parts = []
        for page in doc:
            page_text = page.get_text()
            if page_text:
                text_parts.append(page_text)
        doc.close()
        text = "\n\n".join(text_parts)

        if text.strip() and not is_garbled_text(text):
            return text
        print("PyMuPDF text looks garbled, trying OCR...")
    except ImportError:
        print("PyMuPDF not available, trying OCR...")
    except Exception as e:
        print(f"PyMuPDF failed: {e}, trying OCR...")

    # 2. Try OCR via tesseract + pdf2image
    try:
        from pdf2image import convert_from_path
        import pytesseract

        images = convert_from_path(str(filepath), dpi=300)
        text_parts = []
        for img in images:
            page_text = pytesseract.image_to_string(img)
            if page_text:
                text_parts.append(page_text)
        text = "\n\n".join(text_parts)

        if text.strip():
            return text
        print("OCR produced no text, falling back to pypdf...")
    except ImportError:
        print("pytesseract/pdf2image not available, falling back to pypdf...")
    except Exception as e:
        print(f"OCR failed: {e}, falling back to pypdf...")

    # 3. Fallback to pypdf
    try:
        from pypdf import PdfReader
        reader = PdfReader(str(filepath))
        text_parts = []
        for page in reader.pages:
            page_text = page.extract_text()
            if page_text:
                text_parts.append(page_text)
        text = "\n\n".join(text_parts)
    except Exception as e:
        raise Exception(f"All PDF extraction methods failed: {e}")

    return text


async def extract_text_from_file(filepath: Path, file_type: str) -> str:
    """Extract text content from various file types."""
    text = ""
    
    try:
        if file_type in [".txt", ".md"]:
            async with aiofiles.open(filepath, "r", encoding="utf-8", errors="ignore") as f:
                text = await f.read()
        
        elif file_type == ".pdf":
            text = extract_pdf_text(filepath)
        
        elif file_type in [".docx", ".doc"]:
            try:
                from docx import Document as DocxDocument
                doc = DocxDocument(str(filepath))
                text_parts = []
                for para in doc.paragraphs:
                    if para.text.strip():
                        text_parts.append(para.text)
                text = "\n\n".join(text_parts)
            except Exception as e:
                raise Exception(f"DOCX extraction failed: {e}")
        
        elif file_type == ".csv":
            async with aiofiles.open(filepath, "r", encoding="utf-8", errors="ignore") as f:
                text = await f.read()
        
        elif file_type == ".json":
            import json
            async with aiofiles.open(filepath, "r", encoding="utf-8") as f:
                content = await f.read()
                data = json.loads(content)
                text = json.dumps(data, indent=2)
        
    except Exception as e:
        raise Exception(f"Text extraction failed: {e}")
    
    # Clean up spaced-out text from PDFs (e.g., "I n v o i c e" -> "Invoice")
    if text and file_type == ".pdf":
        text = fix_spaced_text(text)

    return text


def fix_spaced_text(text: str) -> str:
    """Fix text where characters are separated by spaces (common in some PDFs).
    Detects patterns like 'I n v o i c e' and joins them into 'Invoice'."""
    lines = text.split("\n")
    fixed_lines = []

    for line in lines:
        # Check if the line looks spaced out: most single chars separated by spaces
        # Pattern: at least 3 single characters separated by single spaces
        stripped = line.strip()
        if not stripped:
            fixed_lines.append(line)
            continue

        # Count single-char tokens vs multi-char tokens
        tokens = stripped.split(" ")
        if len(tokens) > 3:
            single_char_count = sum(1 for t in tokens if len(t) <= 1)
            ratio = single_char_count / len(tokens)

            if ratio > 0.6:
                # This line is spaced out — join characters
                # Remove single spaces between single characters but keep multi-spaces as word breaks
                fixed = re.sub(r'(?<=\S) (?=\S)', '', stripped)
                # Re-add spaces at likely word boundaries (transitions from lowercase to uppercase, etc.)
                # Also split on double+ spaces that were in the original
                parts = re.split(r'  +', stripped)
                fixed = " ".join(re.sub(r'(?<=\S) (?=\S)', '', p) for p in parts)
                fixed_lines.append(fixed)
                continue

        fixed_lines.append(line)

    return "\n".join(fixed_lines)


async def process_document(document_id: str, filepath: Path, file_type: str):
    """Background task to process and index a document."""
    from ..database import async_session
    
    async with async_session() as db:
        try:
            # Get document
            result = await db.execute(
                select(Document).where(Document.id == document_id)
            )
            doc = result.scalar_one_or_none()
            
            if not doc:
                return
            
            # Update status
            doc.status = "processing"
            await db.commit()
            
            # Extract text
            text = await extract_text_from_file(filepath, file_type)
            
            if not text or not text.strip():
                doc.status = "error"
                doc.error_message = "No text content could be extracted"
                await db.commit()
                return
            
            # Index with RAG
            rag = get_rag_pipeline()
            chunk_count = await rag.add_document(
                document_id=document_id,
                content=text,
                metadata={
                    "filename": doc.original_name,
                    "file_type": file_type
                }
            )
            
            # Update document
            doc.status = "ready"
            doc.chunk_count = chunk_count
            doc.processed_at = datetime.utcnow()
            doc.embedding_model = settings.embedding_model
            await db.commit()
            
        except Exception as e:
            # Update with error
            try:
                result = await db.execute(
                    select(Document).where(Document.id == document_id)
                )
                doc = result.scalar_one_or_none()
                if doc:
                    doc.status = "error"
                    doc.error_message = str(e)
                    await db.commit()
            except Exception:
                pass


@router.post(
    "/upload",
    response_model=DocumentUploadResponse,
    summary="Upload a document",
    description="Upload a document (PDF, TXT, MD, DOCX, CSV, JSON) for RAG indexing. The document will be processed in the background.",
    responses={
        200: {"description": "Document uploaded successfully"},
        400: {"model": ErrorResponse, "description": "Invalid file or file too large"},
        413: {"model": ErrorResponse, "description": "Payload too large"},
        422: {"model": ErrorResponse, "description": "Validation error"},
        429: {"model": ErrorResponse, "description": "Rate limit exceeded"},
        503: {"model": ErrorResponse, "description": "RAG service unavailable"}
    }
)
async def upload_document(
    background_tasks: BackgroundTasks,
    file: UploadFile = File(...),
    db: AsyncSession = Depends(get_db)
):
    """Upload a document for RAG indexing."""
    
    if not file.filename:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="No filename provided"
        )
    
    file_ext = Path(file.filename).suffix.lower()
    if file_ext not in SUPPORTED_EXTENSIONS:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail=f"Unsupported file type. Supported: {', '.join(SUPPORTED_EXTENSIONS)}"
        )
    
    content = await file.read()
    if len(content) > MAX_FILE_SIZE:
        raise HTTPException(
            status_code=status.HTTP_413_REQUEST_ENTITY_TOO_LARGE,
            detail=f"File too large. Maximum size: {MAX_FILE_SIZE // 1024 // 1024}MB"
        )
    
    if not settings.rag_enabled:
        raise HTTPException(
            status_code=status.HTTP_503_SERVICE_UNAVAILABLE,
            detail="RAG service is not enabled"
        )
    
    # Generate unique filename
    document_id = str(uuid.uuid4())
    safe_filename = f"{document_id}{file_ext}"
    
    # Ensure documents directory exists
    docs_dir = settings.data_dir / "documents"
    docs_dir.mkdir(parents=True, exist_ok=True)
    
    filepath = docs_dir / safe_filename
    
    # Save file
    async with aiofiles.open(filepath, "wb") as f:
        await f.write(content)
    
    # Create database record
    doc = Document(
        id=document_id,
        filename=safe_filename,
        original_name=file.filename,
        file_type=file_ext.lstrip("."),
        file_size=len(content),
        status="pending"
    )
    db.add(doc)
    await db.commit()
    
    # Process in background
    background_tasks.add_task(process_document, document_id, filepath, file_ext)
    
    return DocumentUploadResponse(
        id=document_id,
        filename=file.filename,
        status="pending",
        message="Document uploaded and queued for processing"
    )


@router.get(
    "/",
    response_model=PaginatedResponse[DocumentResponse],
    summary="List documents",
    description="Retrieve a paginated list of uploaded documents. Filter by status (pending, processing, ready, error).",
    responses={
        200: {"description": "Successfully retrieved documents"},
        422: {"model": ErrorResponse, "description": "Invalid pagination parameters"}
    }
)
async def list_documents(
    page: int = Query(1, ge=1, description="Page number (1-based)"),
    page_size: int = Query(DEFAULT_PAGE_SIZE, ge=1, le=MAX_PAGE_SIZE, description="Number of items per page"),
    status_filter: str = Query(None, alias="status", description="Filter by document status"),
    db: AsyncSession = Depends(get_db)
):
    """List all uploaded documents with pagination."""
    skip = (page - 1) * page_size
    
    base_query = select(Document)
    count_query = select(func.count(Document.id))
    
    if status_filter:
        base_query = base_query.where(Document.status == status_filter)
        count_query = count_query.where(Document.status == status_filter)
    
    total_result = await db.execute(count_query)
    total = total_result.scalar() or 0
    
    query = base_query.order_by(Document.created_at.desc()).offset(skip).limit(page_size)
    result = await db.execute(query)
    documents = result.scalars().all()
    
    items = [
        DocumentResponse(
            id=d.id,
            filename=d.filename,
            original_name=d.original_name,
            file_type=d.file_type,
            file_size=d.file_size,
            status=d.status,
            error_message=d.error_message,
            chunk_count=d.chunk_count,
            created_at=d.created_at,
            processed_at=d.processed_at
        )
        for d in documents
    ]
    
    total_pages = (total + page_size - 1) // page_size if total > 0 else 1
    
    return PaginatedResponse(
        items=items,
        total=total,
        page=page,
        page_size=page_size,
        total_pages=total_pages
    )


@router.get(
    "/{document_id}",
    response_model=DocumentResponse,
    summary="Get a document",
    description="Retrieve details of a specific uploaded document.",
    responses={
        200: {"description": "Document found and returned"},
        404: {"model": ErrorResponse, "description": "Document not found"}
    }
)
async def get_document(
    document_id: str,
    db: AsyncSession = Depends(get_db)
):
    """Get a specific document."""
    result = await db.execute(
        select(Document).where(Document.id == document_id)
    )
    doc = result.scalar_one_or_none()
    
    if not doc:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Document not found"
        )
    
    return DocumentResponse(
        id=doc.id,
        filename=doc.filename,
        original_name=doc.original_name,
        file_type=doc.file_type,
        file_size=doc.file_size,
        status=doc.status,
        error_message=doc.error_message,
        chunk_count=doc.chunk_count,
        created_at=doc.created_at,
        processed_at=doc.processed_at
    )


@router.delete(
    "/{document_id}",
    status_code=status.HTTP_204_NO_CONTENT,
    summary="Delete a document",
    description="Delete a document and remove it from the RAG index.",
    responses={
        204: {"description": "Document deleted successfully"},
        404: {"model": ErrorResponse, "description": "Document not found"},
        503: {"model": ErrorResponse, "description": "RAG service unavailable"}
    }
)
async def delete_document(
    document_id: str,
    db: AsyncSession = Depends(get_db)
):
    """Delete a document and its index."""
    result = await db.execute(
        select(Document).where(Document.id == document_id)
    )
    doc = result.scalar_one_or_none()
    
    if not doc:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Document not found"
        )
    
    rag = get_rag_pipeline()
    await rag.delete_document(document_id)
    
    filepath = settings.data_dir / "documents" / doc.filename
    if filepath.exists():
        os.remove(filepath)
    
    await db.delete(doc)
    await db.commit()


@router.post(
    "/{document_id}/reprocess",
    summary="Reprocess a document",
    description="Re-index a document that has already been uploaded.",
    responses={
        200: {"description": "Document queued for reprocessing"},
        404: {"model": ErrorResponse, "description": "Document not found"},
        503: {"model": ErrorResponse, "description": "RAG service unavailable"}
    }
)
async def reprocess_document(
    document_id: str,
    background_tasks: BackgroundTasks,
    db: AsyncSession = Depends(get_db)
):
    """Reprocess a document."""
    result = await db.execute(
        select(Document).where(Document.id == document_id)
    )
    doc = result.scalar_one_or_none()
    
    if not doc:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Document not found"
        )
    
    # Delete existing index
    rag = get_rag_pipeline()
    await rag.delete_document(document_id)
    
    # Reset status
    doc.status = "pending"
    doc.error_message = None
    doc.chunk_count = 0
    doc.processed_at = None
    await db.commit()
    
    # Reprocess
    filepath = settings.data_dir / "documents" / doc.filename
    file_ext = Path(doc.filename).suffix.lower()
    
    background_tasks.add_task(process_document, document_id, filepath, file_ext)
    
    return {"message": "Document queued for reprocessing"}
